"""Document parser for extracting text and images from multi-format documents."""

import base64
import logging
from io import BytesIO
from typing import Any, Dict, Optional

from docx import Document
from PIL import Image
from PyPDF2 import PdfReader

from ..exceptions import (
    Base64DecodingError,
    DOCXParsingError,
    ImageParsingError,
    PDFParsingError,
)


log = logging.getLogger(__name__)


class DocumentContext:
    """Shared document context to avoid repeated decoding and metadata extraction."""

    def __init__(self, file_type: str, base64_data: str, raw_bytes: Optional[bytes] = None):
        self.file_type = file_type.lower().strip()
        self.base64_data = base64_data
        self._raw_bytes = raw_bytes

    @property
    def raw_bytes(self) -> bytes:
        if self._raw_bytes is None:
            try:
                self._raw_bytes = base64.b64decode(self.base64_data)
            except base64.binascii.Error as exc:  # pragma: no cover - defensive
                raise Base64DecodingError(f"Invalid base64 encoding: {exc}") from exc
        return self._raw_bytes


class DocumentParser:
    """Parser for extracting text and image content from documents."""
    
    def parse_pdf(
        self,
        context: DocumentContext,
        all_pages: bool = True,
    ) -> str:
        """Extract text from a PDF document.
        
        Args:
            document_base64: Base64 encoded PDF document
            all_pages: If True, extract from all pages; if False, first page only
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If document cannot be decoded or parsed
        """
        try:
            pdf_file = BytesIO(context.raw_bytes)
            reader = PdfReader(pdf_file)
            
            if len(reader.pages) == 0:
                raise PDFParsingError("PDF document has no pages")
            
            # Extract text from pages
            if all_pages:
                # Multi-page extraction
                texts = []
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        texts.append(f"=== Page {page_num} ===\n{page_text.strip()}")
                
                if not texts:
                    raise PDFParsingError("No text could be extracted from any PDF page")
                
                return "\n\n".join(texts)
            else:
                # Single page extraction (backward compatible)
                first_page = reader.pages[0]
                text = first_page.extract_text()
                
                if not text or not text.strip():
                    raise PDFParsingError("No text could be extracted from PDF first page")
                
                return text.strip()
            
        except Base64DecodingError:
            raise
        except PDFParsingError:
            raise
        except Exception as exc:
            raise PDFParsingError(f"Failed to parse PDF document: {exc}") from exc
    
    def parse_docx(
        self,
        context: DocumentContext,
    ) -> str:
        """Extract text from a DOCX document.
        
        Args:
            document_base64: Base64 encoded DOCX document
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If document cannot be decoded or parsed
        """
        try:
            docx_file = BytesIO(context.raw_bytes)
            doc = Document(docx_file)
            
            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        paragraphs.append(" | ".join(row_data))
            
            if not paragraphs:
                raise DOCXParsingError("No text could be extracted from DOCX document")
            
            return "\n\n".join(paragraphs)
            
        except Base64DecodingError:
            raise
        except DOCXParsingError:
            raise
        except Exception as exc:
            raise DOCXParsingError(f"Failed to parse DOCX document: {exc}") from exc
    
    def parse_image(
        self,
        context: DocumentContext,
    ) -> Dict[str, Any]:
        """Parse image and return metadata for vision-based extraction.
        
        Args:
            document_base64: Base64 encoded image
            file_type: Image file type (png, jpg, jpeg)
            
        Returns:
            Dictionary with image data and metadata
            
        Raises:
            ValueError: If image cannot be decoded or parsed
        """
        try:
            image = Image.open(BytesIO(context.raw_bytes))
            
            # Get image metadata
            width, height = image.size
            mode = image.mode
            format_name = image.format or context.file_type.upper()
            
            # Convert to RGB if needed (for consistency)
            if mode not in ['RGB', 'L']:
                image = image.convert('RGB')
            
            # Return image data for vision API
            return {
                "base64_data": context.base64_data,
                "width": width,
                "height": height,
                "mode": mode,
                "format": format_name,
                "media_type": f"image/{context.file_type.lower()}"
            }
            
        except Base64DecodingError:
            raise
        except ImageParsingError:
            raise
        except Exception as exc:
            raise ImageParsingError(f"Failed to parse image: {exc}") from exc


_PARSER = DocumentParser()


def parse_document(
    context: DocumentContext,
    all_pages: bool = True,
) -> str:
    """Parse document and extract text content.
    
    Args:
        context: Shared document context
        all_pages: For PDFs, extract all pages (True) or first page only (False)
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type not supported or parsing fails
    """
    if context.file_type == 'pdf':
        return _PARSER.parse_pdf(context, all_pages)
    elif context.file_type == 'docx':
        return _PARSER.parse_docx(context)
    elif context.file_type in ['png', 'jpg', 'jpeg']:
        from ..exceptions import DocumentParsingError
        raise DocumentParsingError(
            f"Image files ({context.file_type}) require vision-based extraction. "
            "Use parse_image() instead."
        )
    else:
        from ..exceptions import UnsupportedFileTypeError
        raise UnsupportedFileTypeError(
            context.file_type,
            ['pdf', 'docx', 'png', 'jpg', 'jpeg']
        )


def parse_image_document(
    context: DocumentContext,
) -> Dict[str, Any]:
    """Parse image document for vision-based extraction.
    
    Args:
        context: Shared document context for an image
        
    Returns:
        Dictionary with image data and metadata
        
    Raises:
        ValueError: If parsing fails
    """
    return _PARSER.parse_image(context)
